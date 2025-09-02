import streamlit as st
import os
import io
import pandas as pd
import plotly.io as pio
import plotly.express as px
import plotly.graph_objects as go
from datetime import date
import traceback

# --- Dépendances critiques ---
try:
    from fpdf import FPDF
except ImportError:
    FPDF = None

# --- Dépendances pour Word ---
try:
    import docx
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# --- Fonctions de base du projet ---
from .patrimoine_logic import get_patrimoine_df, find_associated_loans
from .charts import (
    create_patrimoine_brut_treemap,
    create_patrimoine_net_donut,
    create_gantt_chart_fig
)

# --- Classe PDF personnalisée pour le rapport ---
class PDF(FPDF):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.font_family_name = "Arial"
        self.currency_symbol = " EUR"

        font_dir = "assets/fonts"
        dejavu_sans_path = os.path.join(font_dir, 'DejaVuSans.ttf')
        dejavu_bold_path = os.path.join(font_dir, 'DejaVuSans-Bold.ttf')
        dejavu_italic_path = os.path.join(font_dir, 'DejaVuSans-Oblique.ttf')

        try:
            self.add_font('DejaVu', '', dejavu_sans_path, uni=True)
            self.add_font('DejaVu', 'B', dejavu_bold_path, uni=True)
            self.add_font('DejaVu', 'I', dejavu_italic_path, uni=True)
            self.font_family_name = "DejaVu"
            self.currency_symbol = " €"
        except (RuntimeError, FileNotFoundError):
            st.error(
                "**Police DejaVu non trouvée.** Le rapport PDF ne peut pas être généré.\n\n"
                "Veuillez créer un dossier `assets/fonts` et y placer les fichiers `DejaVuSans.ttf`, `DejaVuSans-Bold.ttf`, et `DejaVuSans-Oblique.ttf`."
            )
            st.stop()

    def header(self):
        self.set_font(self.font_family_name, 'B', 15)
        self.cell(80)
        self.cell(30, 10, "Rapport d'Audit Patrimonial", 0, 0, 'C')
        self.ln(20)

    def footer(self):
        self.set_y(-15)
        self.set_font(self.font_family_name, 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', 0, 0, 'C')

    def chapter_title(self, title):
        self.set_font(self.font_family_name, 'B', 14)
        self.set_fill_color(230, 230, 230)
        self.cell(0, 10, title, 0, 1, 'L', 1)
        self.ln(5)

    def chapter_body(self, body):
        self.set_font(self.font_family_name, '', 11)
        self.multi_cell(0, 6, body)
        self.ln()

    def add_table(self, df, col_widths, title=""):
        if title:
            self.set_font(self.font_family_name, 'B', 12)
            self.cell(0, 10, title, 0, 1, 'L')
        self.set_font(self.font_family_name, 'B', 10)
        self.set_fill_color(240, 240, 240)
        for i, col_name in enumerate(df.columns):
            self.cell(col_widths[i], 7, str(col_name), 1, 0, 'C', 1)
        self.ln()
        self.set_font(self.font_family_name, '', 9)
        for _, row in df.iterrows():
            for i, item in enumerate(row):
                text = str(item)
                align = 'L'
                if isinstance(item, (int, float)):
                    text = f"{item:,.2f}{self.currency_symbol}"
                    align = 'R'
                self.cell(col_widths[i], 6, text, 1, 0, align)
            self.ln()
        self.ln(10)

    def add_figure_to_pdf(self, fig, width_percent=100, title=""):
        if title:
            self.set_font(self.font_family_name, 'B', 12)
            self.cell(0, 10, title, 0, 1, 'L')
        try:
            page_width = self.w - 2 * self.l_margin
            img_width = page_width * (width_percent / 100)
            img_height = img_width * 0.6
            img_bytes = pio.to_image(fig, format="png", width=int(img_width*5), height=int(img_height*5), scale=2)
            
            x_pos = (self.w - img_width) / 2
            self.image(io.BytesIO(img_bytes), x=x_pos, w=img_width)
            self.ln(img_height + 5)
        except Exception as e:
            self.set_text_color(255, 0, 0)
            self.multi_cell(0, 5, f"Erreur lors de la génération du graphique : {e}")
            self.set_text_color(0, 0, 0)
            self.ln()

# --- Fonctions pour générer les sections du rapport ---

def add_title_page(pdf, parents):
    pdf.add_page()
    pdf.set_font(pdf.font_family_name, 'B', 24)
    pdf.cell(0, 80, "Audit Patrimonial", 0, 1, 'C')
    pdf.set_font(pdf.font_family_name, 'I', 16)
    if parents:
        parent_names = " & ".join([p['prenom'] for p in parents if p.get('prenom')])
        pdf.cell(0, 20, f"Préparé pour {parent_names}", 0, 1, 'C')
    pdf.set_font(pdf.font_family_name, '', 12)
    pdf.cell(0, 10, f"Date du rapport : {date.today().strftime('%d/%m/%Y')}", 0, 1, 'C')

def add_family_section(pdf, parents, enfants):
    pdf.add_page()
    pdf.chapter_title("1. Composition du Foyer")
    pdf.set_font(pdf.font_family_name, 'B', 12)
    pdf.cell(0, 10, "Parents", 0, 1, 'L')
    pdf.set_font(pdf.font_family_name, '', 11)
    for p in parents:
        dob_str = p.get('date_naissance', date.today()).strftime('%d/%m/%Y')
        pdf.cell(0, 6, f"- {p.get('prenom', 'N/A')}, né(e) le {dob_str}", 0, 1, 'L')
    pdf.ln(5)
    if enfants:
        pdf.set_font(pdf.font_family_name, 'B', 12)
        pdf.cell(0, 10, "Enfants", 0, 1, 'L')
        pdf.set_font(pdf.font_family_name, '', 11)
        for e in enfants:
            dob_str = e.get('date_naissance', date.today()).strftime('%d/%m/%Y')
            pdf.cell(0, 6, f"- {e.get('prenom', 'N/A')}, né(e) le {dob_str}", 0, 1, 'L')

def add_patrimoine_section(pdf, actifs, passifs):
    pdf.add_page()
    pdf.chapter_title("2. Bilan et Visualisation du Patrimoine")
    df_patrimoine = get_patrimoine_df(actifs, passifs)
    if df_patrimoine.empty:
        pdf.chapter_body("Aucune donnée patrimoniale n'a été renseignée.")
        return
    total_actifs = df_patrimoine['Valeur Brute'].sum()
    total_passifs = df_patrimoine['Passif'].sum()
    patrimoine_net = df_patrimoine['Valeur Nette'].sum()
    pdf.chapter_body(
        f"Le total de vos actifs s'élève à {total_actifs:,.2f}{pdf.currency_symbol}, pour un passif de {total_passifs:,.2f}{pdf.currency_symbol}.\n"
        f"Votre patrimoine net est donc de {patrimoine_net:,.2f}{pdf.currency_symbol}."
    )
    df_display = df_patrimoine[['Libellé', 'Type', 'Valeur Brute', 'Passif', 'Valeur Nette']]
    pdf.add_table(df_display, [60, 40, 30, 30, 30], "Détail du patrimoine")
    
    fig_brut = create_patrimoine_brut_treemap(df_patrimoine)
    pdf.add_figure_to_pdf(fig_brut, width_percent=80, title="Répartition du Patrimoine Brut")
    
    fig_net_donut = create_patrimoine_net_donut(df_patrimoine)
    pdf.add_figure_to_pdf(fig_net_donut, width_percent=80, title="Répartition du Patrimoine Net")

def add_flux_section(pdf, revenus, depenses):
    pdf.add_page()
    pdf.chapter_title("3. Flux Financiers Mensuels")
    total_revenus = sum(r.get('montant', 0.0) for r in revenus)
    total_depenses = sum(d.get('montant', 0.0) for d in depenses)
    capacite_epargne = total_revenus - total_depenses
    pdf.chapter_body(
        f"Le total de vos revenus mensuels est de {total_revenus:,.2f}{pdf.currency_symbol} pour des dépenses de {total_depenses:,.2f}{pdf.currency_symbol}.\n"
        f"Votre capacité d'épargne mensuelle est de {capacite_epargne:,.2f}{pdf.currency_symbol}."
    )
    if revenus:
        df_revenus = pd.DataFrame(revenus)[['libelle', 'montant', 'type']]
        df_revenus.columns = ['Libellé', 'Montant', 'Type']
        pdf.add_table(df_revenus, [90, 50, 50], "Détail des revenus")
    if depenses:
        df_depenses = pd.DataFrame(depenses)[['libelle', 'montant', 'categorie']]
        df_depenses.columns = ['Libellé', 'Montant', 'Catégorie']
        pdf.add_table(df_depenses, [90, 50, 50], "Détail des dépenses")

def add_capacite_endettement_section(pdf, revenus, passifs, max_debt_ratio, funcs):
    pdf.add_page()
    pdf.chapter_title("4. Capacité d'Endettement")
    if not funcs['CAP_ENDETTEMENT_AVAILABLE']:
        pdf.chapter_body("Le module de calcul de la capacité d'endettement n'est pas disponible.")
        return
        
    weighted_income_data = funcs['calculate_weighted_income'](revenus)
    debt_data = funcs['calculate_current_debt_service'](passifs)
    total_weighted_income = weighted_income_data["total"]
    total_current_debt = debt_data["total"]
    
    if total_weighted_income == 0:
        pdf.chapter_body("Les revenus pondérés sont nuls. Impossible de calculer la capacité d'endettement.")
        return

    current_debt_ratio_pct = (total_current_debt / total_weighted_income) * 100 if total_weighted_income > 0 else 0
    max_debt_service = total_weighted_income * (max_debt_ratio / 100)
    remaining_capacity = max(0, max_debt_service - total_current_debt)

    pdf.chapter_body(
        f"Vos revenus mensuels pondérés sont de {total_weighted_income:,.2f}{pdf.currency_symbol}.\n"
        f"Vos charges de prêts actuelles sont de {total_current_debt:,.2f}{pdf.currency_symbol}.\n"
        f"Votre taux d'endettement actuel est de {current_debt_ratio_pct:.2f} % (pour une cible de {max_debt_ratio} %).\n"
        f"Votre capacité de remboursement mensuelle restante est de {remaining_capacity:,.2f}{pdf.currency_symbol}."
    )
    
    fig_gauge = go.Figure(go.Indicator(
        mode="gauge+number", value=current_debt_ratio_pct, number={'suffix': ' %'},
        title={'text': f"Taux d'endettement (Cible: {max_debt_ratio}%)"},
        gauge={'axis': {'range': [None, 50]}, 'bar': {'color': "darkblue"},
               'steps': [{'range': [0, max_debt_ratio], 'color': 'lightgreen'}, {'range': [max_debt_ratio, 50], 'color': 'lightcoral'}],
               'threshold': {'line': {'color': "red", 'width': 4}, 'thickness': 0.75, 'value': max_debt_ratio}}))
    pdf.add_figure_to_pdf(fig_gauge, width_percent=70)

def add_projection_section(pdf, parents, enfants, passifs, settings, duree_projection, funcs):
    pdf.add_page()
    pdf.chapter_title("5. Projection des Étapes de Vie")
    if not funcs['PROJECTION_AVAILABLE']:
        pdf.chapter_body("Le module de projection n'est pas disponible.")
        return

    gantt_data = funcs['generate_gantt_data'](parents, enfants, settings, duree_projection)
    df_projection = funcs['generate_financial_projection'](parents, enfants, passifs, settings, duree_projection)

    if not gantt_data or df_projection.empty:
        pdf.chapter_body("Données insuffisantes pour générer les projections.")
        return

    fig_gantt = create_gantt_chart_fig(gantt_data, duree_projection, parents, enfants)
    if fig_gantt:
        fig_gantt.update_layout(
            title_text='Frise Chronologique du Foyer avec âges aux étapes clés',
            height=max(300, len(parents + enfants) * 60 + 100)
        )
        pdf.add_figure_to_pdf(fig_gantt, width_percent=100)

    cols_to_stack = [col for col in ['Reste à vivre', 'Impôt sur le revenu', 'Prélèvements Sociaux', 'Mensualités Prêts', 'Autres Dépenses', 'Coût des études'] if col in df_projection.columns]
    fig_proj_fin = px.bar(df_projection, x='Année', y=cols_to_stack, title="Répartition des revenus du foyer", barmode='stack')
    if 'Revenus du foyer' in df_projection.columns:
        fig_proj_fin.add_scatter(x=df_projection['Année'], y=df_projection['Revenus du foyer'], mode='lines', name='Total Revenus', line=dict(color='black', dash='dot'))
    pdf.add_figure_to_pdf(fig_proj_fin, width_percent=100)

def add_immo_focus_section(pdf, actifs, passifs, tmi, projection_duration, funcs):
    productive_assets = [a for a in actifs if a.get('type') == "Immobilier productif"]
    if not productive_assets: return

    pdf.add_page()
    pdf.chapter_title("6. Focus Immobilier Locatif")
    if not funcs['FOCUS_IMMO_AVAILABLE']:
        pdf.chapter_body("Le module d'analyse immobilière n'est pas disponible.")
        return

    for asset in productive_assets:
        if asset.get('loyers_mensuels') is None: continue
        pdf.set_font(pdf.font_family_name, 'B', 12)
        pdf.cell(0, 10, f"Analyse de : {asset.get('libelle', 'Sans nom')}", 0, 1, 'L')
        pdf.ln(2)
        loans = find_associated_loans(asset.get('id'), passifs)
        df_projection = funcs['generate_immo_projection_data'](asset, loans, tmi, 17.2, projection_duration)
        if df_projection.empty: continue
        
        fig_cf = funcs['create_cash_flow_projection_fig'](df_projection)
        pdf.add_figure_to_pdf(fig_cf, width_percent=80)

def add_fiscalite_section(pdf, parents, enfants, revenus_salaires, revenu_foncier_net, annee, revenu_max_graphique, est_parent_isole, funcs):
    pdf.add_page()
    pdf.chapter_title("7. Focus Fiscalité")
    if not funcs['OPENFISCA_AVAILABLE'] or not funcs['FOCUS_FISCALITE_AVAILABLE']:
        pdf.chapter_body("Le module d'analyse fiscale (OpenFisca) n'est pas disponible.")
        return

    results = funcs['analyser_fiscalite_foyer_fiscalite'](annee, parents, enfants, revenus_salaires, revenu_foncier_net, est_parent_isole)
    
    pdf.chapter_body(
        f"Pour l'année {annee}, votre impôt sur le revenu net est estimé à {results['ir_net']:,.2f}{pdf.currency_symbol}.\n"
        f"Votre Taux Marginal d'Imposition (TMI) est de {results['tmi']:.0f} %.\n"
        f"L'économie d'impôt liée au quotient familial est de {results['gain_quotient']:,.2f}{pdf.currency_symbol}."
    )
    
    df_evol, bareme = funcs['simuler_evolution_fiscalite'](annee, parents, enfants, revenu_foncier_net, est_parent_isole, revenu_max_graphique)
    results['df_evolution'] = df_evol
    results['bareme'] = bareme
    
    fig = funcs['display_income_evolution_chart'](results, sum(revenus_salaires.values()), revenu_max_graphique)
    if fig:
        pdf.add_figure_to_pdf(fig, width_percent=100)

def add_per_section(pdf, per_results, funcs):
    pdf.add_page()
    pdf.chapter_title("8. Optimisation PER")
    if not funcs['OPENFISCA_AVAILABLE'] or not funcs['OPTIMISATION_PER_AVAILABLE']:
        pdf.chapter_body("Le module d'optimisation PER (OpenFisca) n'est pas disponible.")
        return
    if not per_results or "error" in per_results:
        pdf.chapter_body("Aucune simulation PER n'a été effectuée ou la simulation a échoué.")
        return

    ir_initial = per_results['df_one_shot']['IR'].values[0]
    versement_optimal = per_results['versement_optimal_PER']
    impot_final = per_results['impot_avec_versement']
    economie = ir_initial - impot_final
    
    pdf.chapter_body(
        f"Votre impôt initial est de {ir_initial:,.2f}{pdf.currency_symbol}.\n"
        f"Un versement PER optimal de {versement_optimal:,.2f}{pdf.currency_symbol} permettrait de réduire votre impôt à {impot_final:,.2f}{pdf.currency_symbol}.\n"
        f"L'économie d'impôt réalisée serait de {economie:,.2f}{pdf.currency_symbol}."
    )

def generate_report_safe(selections, parents, enfants, actifs, passifs, revenus, depenses, per_results, settings, funcs):
    """
    Génère le rapport PDF et gère les erreurs.
    Retourne un tuple (pdf_bytes, error_message).
    """
    if FPDF is None:
        return None, "La bibliothèque `fpdf2` est nécessaire. Installez-la avec `pip install fpdf2`."
    try:
        pdf = PDF()
        pdf.alias_nb_pages()
        add_title_page(pdf, parents)
        if selections['famille']: add_family_section(pdf, parents, enfants)
        if selections['patrimoine']: add_patrimoine_section(pdf, actifs, passifs)
        if selections['flux']: add_flux_section(pdf, revenus, depenses)
        if selections['endettement']: add_capacite_endettement_section(pdf, revenus, passifs, settings['max_debt_ratio'], funcs)
        if selections['projection']: add_projection_section(pdf, parents, enfants, passifs, settings['projection_settings'], settings['duree_projection'], funcs)
        if selections['immo']: add_immo_focus_section(pdf, actifs, passifs, settings['immo_tmi'], settings['immo_projection_duration'], funcs)
        if (selections['fiscalite'] or selections['per']) and funcs.get('get_revenus_imposables_fiscalite'):
            revenus_salaires, revenu_foncier_net = funcs['get_revenus_imposables_fiscalite'](settings['annee_fiscalite'])
            if selections['fiscalite']: add_fiscalite_section(pdf, parents, enfants, revenus_salaires, revenu_foncier_net, settings['annee_fiscalite'], settings['revenu_max_fiscalite'], settings['est_parent_isole_fiscalite'], funcs)
            if selections['per']: add_per_section(pdf, per_results, funcs)
        
        return bytes(pdf.output(dest='S')), None
    except Exception as e:
        return None, f"Erreur lors de la génération du rapport : {e}\n{traceback.format_exc()}"

# --- NOUVELLES FONCTIONS POUR LA GÉNÉRATION WORD ---

def add_word_figure(document, fig, width_inches=6.0):
    """Ajoute une figure Plotly à un document Word."""
    try:
        img_bytes = pio.to_image(fig, format="png", scale=2, width=int(width_inches*300), height=int(width_inches*0.6*300))
        document.add_picture(io.BytesIO(img_bytes), width=Inches(width_inches))
    except Exception as e:
        document.add_paragraph(f"Erreur lors de la génération du graphique : {e}").runs[0].font.color.rgb = docx.shared.RGBColor(255, 0, 0)

def add_word_table(document, df, title=""):
    """Ajoute un DataFrame pandas comme tableau dans un document Word."""
    if title:
        document.add_heading(title, level=2)
    
    table = document.add_table(rows=1, cols=df.shape[1])
    table.style = 'Table Grid'
    
    # En-têtes
    for i, col_name in enumerate(df.columns):
        table.cell(0, i).text = str(col_name)

    # Données
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, item in enumerate(row):
            text = str(item)
            if isinstance(item, (int, float)):
                text = f"{item:,.2f} €"
            cells[i].text = text

def generate_word_report_safe(selections, parents, enfants, actifs, passifs, revenus, depenses, per_results, settings, funcs):
    """
    Génère le rapport Word et gère les erreurs.
    Retourne un tuple (word_bytes, error_message).
    """
    if not DOCX_AVAILABLE:
        return None, "La bibliothèque `python-docx` est nécessaire. Installez-la avec `pip install python-docx`."

    try:
        document = docx.Document()
        
        # --- Page de titre ---
        document.add_heading("Audit Patrimonial", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        if parents:
            parent_names = " & ".join([p['prenom'] for p in parents if p.get('prenom')])
            p = document.add_paragraph(f"Préparé pour {parent_names}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph(f"Date du rapport : {date.today().strftime('%d/%m/%Y')}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_page_break()

        # --- Sections ---
        if selections['famille']:
            document.add_heading("1. Composition du Foyer", level=1)
            document.add_heading("Parents", level=2)
            for p in parents:
                dob_str = p.get('date_naissance', date.today()).strftime('%d/%m/%Y')
                document.add_paragraph(f"- {p.get('prenom', 'N/A')}, né(e) le {dob_str}", style='List Bullet')
            if enfants:
                document.add_heading("Enfants", level=2)
                for e in enfants:
                    dob_str = e.get('date_naissance', date.today()).strftime('%d/%m/%Y')
                    document.add_paragraph(f"- {e.get('prenom', 'N/A')}, né(e) le {dob_str}", style='List Bullet')

        if selections['patrimoine']:
            document.add_heading("2. Bilan et Visualisation du Patrimoine", level=1)
            df_patrimoine = get_patrimoine_df(actifs, passifs)
            if not df_patrimoine.empty:
                total_actifs = df_patrimoine['Valeur Brute'].sum()
                total_passifs = df_patrimoine['Passif'].sum()
                patrimoine_net = df_patrimoine['Valeur Nette'].sum()
                document.add_paragraph(f"Votre patrimoine net est de {patrimoine_net:,.2f} €.")
                df_display = df_patrimoine[['Libellé', 'Type', 'Valeur Brute', 'Passif', 'Valeur Nette']]
                add_word_table(document, df_display, "Détail du patrimoine")
                
                document.add_heading("Répartition du Patrimoine Brut", level=2)
                fig_brut = create_patrimoine_brut_treemap(df_patrimoine)
                add_word_figure(document, fig_brut)

        if selections['flux']:
            document.add_heading("3. Flux Financiers Mensuels", level=1)
            total_revenus = sum(r.get('montant', 0.0) for r in revenus)
            total_depenses = sum(d.get('montant', 0.0) for d in depenses)
            capacite_epargne = total_revenus - total_depenses
            document.add_paragraph(f"Votre capacité d'épargne mensuelle est de {capacite_epargne:,.2f} €.")
            if revenus:
                df_revenus = pd.DataFrame(revenus)[['libelle', 'montant', 'type']]
                df_revenus.columns = ['Libellé', 'Montant', 'Type']
                add_word_table(document, df_revenus, "Détail des revenus")

        if selections['endettement']:
            document.add_heading("4. Capacité d'Endettement", level=1)
            if funcs['CAP_ENDETTEMENT_AVAILABLE']:
                weighted_income_data = funcs['calculate_weighted_income'](revenus)
                debt_data = funcs['calculate_current_debt_service'](passifs)
                total_weighted_income = weighted_income_data["total"]
                total_current_debt = debt_data["total"]
                max_debt_ratio = settings['max_debt_ratio']
                if total_weighted_income > 0:
                    current_debt_ratio_pct = (total_current_debt / total_weighted_income) * 100
                    remaining_capacity = max(0, (total_weighted_income * (max_debt_ratio / 100)) - total_current_debt)
                    document.add_paragraph(f"Votre taux d'endettement actuel est de {current_debt_ratio_pct:.2f} % (cible: {max_debt_ratio} %).")
                    document.add_paragraph(f"Votre capacité de remboursement mensuelle restante est de {remaining_capacity:,.2f} €.")

        # ... Ajoutez ici la logique pour les autres sections (projection, immo, etc.) de la même manière ...
        # Par exemple pour la projection :
        if selections['projection'] and funcs['PROJECTION_AVAILABLE']:
            document.add_heading("5. Projection des Étapes de Vie", level=1)
            gantt_data = funcs['generate_gantt_data'](parents, enfants, settings['projection_settings'], settings['duree_projection'])
            fig_gantt = create_gantt_chart_fig(gantt_data, settings['duree_projection'], parents, enfants)
            if fig_gantt:
                add_word_figure(document, fig_gantt)

        # --- Section Focus Immobilier ---
        if selections['immo'] and funcs.get('FOCUS_IMMO_AVAILABLE'):
            document.add_heading("6. Focus Immobilier Locatif", level=1)
            productive_assets = [a for a in actifs if a.get('type') == "Immobilier productif"]
            if not productive_assets:
                document.add_paragraph("Aucun bien immobilier productif n'a été renseigné.")
            else:
                for asset in productive_assets:
                    if asset.get('loyers_mensuels') is None: continue
                    document.add_heading(f"Analyse de : {asset.get('libelle', 'Sans nom')}", level=2)
                    
                    metrics = funcs['calculate_property_metrics'](asset, passifs, settings['immo_tmi'], 17.2, date.today().year)
                    
                    p = document.add_paragraph()
                    p.add_run('Rentabilité Brute: ').bold = True
                    p.add_run(f"{metrics['gross_yield']:.2f} % | ")
                    p.add_run('Nette de charges: ').bold = True
                    p.add_run(f"{metrics['net_yield_charges']:.2f} % | ")
                    p.add_run('Nette de fiscalité: ').bold = True
                    p.add_run(f"{metrics['net_yield_after_tax']:.2f} %")
                    
                    p = document.add_paragraph()
                    p.add_run("Effort d'épargne mensuel: ").bold = True
                    p.add_run(f"{metrics['savings_effort']:,.2f} €")

                    is_lmnp = asset.get('mode_exploitation') == 'Location Meublée'
                    fig_waterfall = funcs['_create_waterfall_fig'](metrics, date.today().year, is_lmnp=is_lmnp)
                    add_word_figure(document, fig_waterfall)

                    df_projection = funcs['generate_immo_projection_data'](asset, find_associated_loans(asset.get('id'), passifs), settings['immo_tmi'], 17.2, settings['immo_projection_duration'])
                    if not df_projection.empty:
                        document.add_heading(f"Projections sur {settings['immo_projection_duration']} ans", level=3)
                        fig_cf = funcs['create_cash_flow_projection_fig'](df_projection)
                        add_word_figure(document, fig_cf)
                        fig_lev = funcs['create_leverage_projection_fig'](df_projection)
                        add_word_figure(document, fig_lev)

        # --- Section Focus Fiscalité ---
        if selections['fiscalite'] and funcs.get('OPENFISCA_AVAILABLE') and funcs.get('FOCUS_FISCALITE_AVAILABLE'):
            document.add_heading("7. Focus Fiscalité", level=1)
            revenus_salaires, revenu_foncier_net = funcs['get_revenus_imposables_fiscalite'](settings['annee_fiscalite'])
            results = funcs['analyser_fiscalite_foyer_fiscalite'](settings['annee_fiscalite'], parents, enfants, revenus_salaires, revenu_foncier_net, settings['est_parent_isole_fiscalite'])
            
            document.add_paragraph(f"Pour l'année {settings['annee_fiscalite']}, votre impôt sur le revenu net est estimé à {results['ir_net']:,.2f} €.")
            document.add_paragraph(f"Votre Taux Marginal d'Imposition (TMI) est de {results['tmi']:.0f} %.")
            document.add_paragraph(f"L'économie d'impôt liée au quotient familial est de {results['gain_quotient']:,.2f} €.")

            fig = funcs['display_income_evolution_chart'](results, sum(revenus_salaires.values()), settings['revenu_max_fiscalite'])
            if fig:
                add_word_figure(document, fig)

        # --- Section Optimisation PER ---
        if selections['per'] and funcs.get('OPENFISCA_AVAILABLE') and funcs.get('OPTIMISATION_PER_AVAILABLE'):
            document.add_heading("8. Optimisation PER", level=1)
            if not per_results or "error" in per_results:
                document.add_paragraph("Aucune simulation PER n'a été effectuée ou la simulation a échoué.")
            else:
                ir_initial = per_results['df_one_shot']['IR'].values[0]
                versement_optimal = per_results['versement_optimal_PER']
                impot_final = per_results['impot_avec_versement']
                economie = ir_initial - impot_final
                
                document.add_paragraph(f"Votre impôt initial est de {ir_initial:,.2f} €.")
                document.add_paragraph(f"Un versement PER optimal de {versement_optimal:,.2f} € permettrait de réduire votre impôt à {impot_final:,.2f} €.")
                document.add_paragraph(f"L'économie d'impôt réalisée serait de {economie:,.2f} €.")

                if funcs.get('create_base_tax_evolution_fig_per'):
                    fig_per = funcs['create_base_tax_evolution_fig_per'](per_results)
                    fig_per.add_scatter(x=[per_results['total_salary_input']], y=[ir_initial], marker=dict(color='red', size=10), name='IR initial')
                    if versement_optimal > 0:
                        equivalent_row = per_results['df_income_tax_evol'].iloc[(per_results['df_income_tax_evol']['IR'] - impot_final).abs().argmin()]
                        fig_per.add_scatter(x=[equivalent_row['Revenu']], y=[equivalent_row['IR']], marker=dict(color='green', size=10), name='IR après versement')
                    fig_per.update_layout(title="Impact du versement PER")
                    add_word_figure(document, fig_per)

        # Sauvegarde du document dans un flux binaire en mémoire
        doc_io = io.BytesIO()
        document.save(doc_io)
        doc_io.seek(0)
        return doc_io.getvalue(), None

    except Exception as e:
        return None, f"Erreur lors de la génération du rapport Word : {e}\n{traceback.format_exc()}"
